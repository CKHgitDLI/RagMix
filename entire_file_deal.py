import os
from docx import Document
import re
import settings
import shutil
import os
from win32com import client as wc
import subprocess
from pdf2docx import parse


def kill_process_holding_file(filepath):
    # 在Windows系统中
    if os.name == 'nt':
        # 找出并终止占用文件的进程
        cmd = f'handle {filepath} -nobanner -accepteula | findstr "pid:"'
        result = subprocess.run(cmd, capture_output=True, text=True, shell=True)
        if result.stdout:
            pid = result.stdout.strip().split()[-1]
            subprocess.run(f'taskkill /F /PID {pid}', shell=True)
        else:
            print("No locking handle found.")

    # 在Unix-like系统中
    else:
        cmd = f'lsof {filepath}'
        result = subprocess.run(cmd, capture_output=True, text=True, shell=True)
        if result.stdout:
            lines = result.stdout.strip().split('\n')
            if len(lines) > 1:
                pid = lines[1].split()[1]
                subprocess.run(f'kill -9 {pid}', shell=True)
        else:
            print("No process is locking the file.")


def copy_and_rename_file(src_file):
    print(src_file)
    new_file = 'E:\\三库资料\\temp.docx'  # 新的文件名
    print(new_file)
    w = wc.gencache.EnsureDispatch('kwps.application')
    doc = w.Documents.Open(src_file)
    doc.SaveAs2(new_file, 12)  # 问题出在这，必须为12
    doc.Close()
    return new_file


def deal_all_doc_pdf():
    folder_path = 'E:\\三库资料'
    files_docx = []
    files_pdf = []
    files_temp = []
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            # 获取文件的绝对路径
            file_path = os.path.join(root, file)
            print(file_path)
            files_temp.append(file_path)
    for i in files_temp:
        if i.split(".")[-1] != "pdf":
            files_docx.append(i)
        else:
            files_pdf.append(i)
    for i in files_docx:
        print(i)
        copy_and_rename_file(i)
        if i.split(".")[-1] == "doc":
            os.remove(i)
    for i in files_pdf:
        print(i)
        parse(i, new_file=os.path.join(os.path.dirname(i), os.path.basename(i)[0] + '.docx'))  # 新的文件名
        os.remove(i)


def extract_text_from_docx(file_path):
    # if re.search(r"\.doc$", file_path, re.IGNORECASE):
    #     print("正在转换为docx，临时文件在解析后将被自动清理")
    #     filename = settings.convert_doc_to_docx(file_path)
    #     file_path = os.path.join(settings.get_project_base_directory(), filename)
    # 加载Word文档
    doc = Document(file_path)
    # 提取所有段落中的文字
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    # 提取所有表格中的文字
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    # 合并所有文本为一个字符串，可以根据需要添加分隔符
    return '\n'.join(full_text)


def search_files(directory, keyword):
    for root, dirs, files in os.walk(directory):
        for file in files:
            if keyword.lower() in file.lower():
                print(os.path.join(root, file))
                return os.path.join(root, file)


def remove_blank_lines(text):
    return re.sub(r'\n\s*\n', '\n', text)


def entire_file(filename):
    temp = search_files('E:\三库资料2', filename)
    text = extract_text_from_docx(copy_and_rename_file(temp))
    print(remove_blank_lines(text))
    return "文件名：{" + filename + "}\n\n内容：{" + remove_blank_lines(text) + "}\n\n\n\n"


if __name__ == '__main__':
    entire_file("关于梨园河煤矿井下北翼5#皮带故障处理的建议")
