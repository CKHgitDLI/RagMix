import re
from io import BytesIO
from docx import Document
import settings
from parser_file.utils import get_text
from nlp import bullets_category, remove_contents_table, hierarchical_merge, make_colon_as_title, tokenize_chunks, \
    docx_question_level, tokenize_table
from nlp import rag_tokenizer
from parser_file.ragflow_docx import RAGFlowDocxParser as DocxParser
import os
from parser_file.ragflow_pdf import RAGFlowPdfParser as PdfParser
from settings import DEBUG
from parser_file.ragflow_pdf import PlainParser


class Docx(DocxParser):
    def __init__(self):
        pass

    def __clean(self, line):
        line = re.sub(r"\u3000", " ", line).strip()
        return line

    def __call__(self, filename, binary=None, from_page=0, to_page=100000):
        self.doc = Document(
            filename) if not binary else Document(BytesIO(binary))
        tbls = []
        for tb in self.doc.tables:
            # 定位表头
            for i in range(len(self.doc.paragraphs)):
                # 步骤3: 判断每个段落是否为表格
                if self.doc.paragraphs[i]._element.tag.endswith('tbl'):
                    if self.doc.paragraphs[i]._element == tb:
                        print(tb)
            html = "<table>"
            for r in tb.rows:
                html += "<tr>"
                i = 0
                while i < len(r.cells):
                    span = 1
                    c = r.cells[i]
                    for j in range(i + 1, len(r.cells)):
                        if c.text == r.cells[j].text:
                            span += 1
                            i = j
                    i += 1
                    html += f"<td>{c.text}</td>" if span == 1 else f"<td colspan='{span}'>{c.text}</td>"
                html += "</tr>"
            html += "</table>"
            tbls.append(((None, html), ""))
        pn = 0
        lines = []
        bull = bullets_category([p.text for p in self.doc.paragraphs])
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            question_level, p_text = docx_question_level(p, bull)
            if not p_text.strip("\n"): continue
            lines.append((question_level, p_text))

            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1

        visit = [False for _ in range(len(lines))]
        sections = []
        for s in range(len(lines)):
            e = s + 1
            while e < len(lines):
                if lines[e][0] <= lines[s][0]:
                    break
                e += 1
            if e - s == 1 and visit[s]: continue
            sec = []
            next_level = lines[s][0] + 1
            while not sec and next_level < 22:
                for i in range(s + 1, e):
                    if lines[i][0] != next_level: continue
                    sec.append(lines[i][1])
                    visit[i] = True
                next_level += 1
            sec.insert(0, lines[s][1])

            sections.append("\n".join(sec))

        return [l for l in sections if l], tbls

    def __str__(self) -> str:
        return f'''
            question:{self.question},
            answer:{self.answer},
            level:{self.level},
            childs:{self.childs}
        '''


class Pdf(PdfParser):
    def __init__(self):
        self.model_speciess = "laws"
        super().__init__()

    def __call__(self, filename, binary=None, from_page=0,
                 to_page=100000, zoomin=3, callback=None):
        callback("准备启动PDF OCR")
        self.__images__(
            filename if not binary else binary,
            zoomin,
            from_page,
            to_page,
            callback
        )
        callback("PDF OCR结束")

        from timeit import default_timer as timer
        start = timer()
        self._layouts_rec(zoomin)
        callback("布局分析结束")
        if DEBUG: print("布局:".format(
            (timer() - start) / (self.total_page + 0.1)))
        self._naive_vertical_merge()

        callback("文本提取结束")

        return [(b["text"], self._line_tag(b, zoomin))
                for b in self.boxes], None


def chunk(filename, binary=None, from_page=0, to_page=100000,
          lang="Chinese", callback=print, **kwargs):
    """
        Supported file formats are docx, pdf, txt.
    """
    print("开始解析" + filename)
    doc = {
        "docnm_kwd": ".".join(os.path.split(filename)[-1].split(".")[:-1]),
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", os.path.split(filename)[-1]))
    }
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    pdf_parser = None
    sections = []
    # is it English
    eng = lang.lower() == "english"  # is_english(sections)

    if re.search(r"\.docx$", os.path.split(filename)[-1], re.IGNORECASE):
        callback("开始解析")
        txts, tbls = Docx()(filename, binary)
        for txt in txts:
            sections.append(txt)
        res = tokenize_table(tbls, doc, eng)
        callback("解析结束")
        chunks = sections
        res.extend(tokenize_chunks(chunks, doc, eng, pdf_parser))
        return res

    elif re.search(r"\.pdf$", filename, re.IGNORECASE):
        pdf_parser = Pdf() if kwargs.get(
            "parser_config", {}).get(
            "layout_recognize", True) else PlainParser()
        for txt, poss in pdf_parser(filename if not binary else binary,
                                    from_page=from_page, to_page=to_page, callback=callback)[0]:
            sections.append(txt + poss)

    elif re.search(r"\.txt$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        txt = get_text(filename, binary)
        sections = txt.split("\n")
        sections = [l for l in sections if l]
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.doc$", filename, re.IGNORECASE):
        print("正在转换为docx，临时文件在解析后将被自动清理")
        filename = settings.convert_doc_to_docx(filename)
        filename = os.path.join(settings.get_project_base_directory(), filename)
        callback("开始解析")
        txts, tbls = Docx()(filename, binary)
        for txt in txts:
            sections.append(txt)
        res = tokenize_table(tbls, doc, eng)
        callback("解析结束")
        chunks = sections
        res.extend(tokenize_chunks(chunks, doc, eng, pdf_parser))
        return res

    else:
        raise NotImplementedError(
            "file type not supported yet(doc, docx, pdf, txt supported)")

    # Remove 'Contents' part
    remove_contents_table(sections, eng)

    make_colon_as_title(sections)
    bull = bullets_category(sections)
    chunks = hierarchical_merge(bull, sections, 5)
    if not chunks:
        callback(0.99, "No chunk parsed out.")

    return tokenize_chunks(["\n".join(ck)
                            for ck in chunks], doc, eng, pdf_parser)


if __name__ == "__main__":
    b = chunk(r"E:\Rag-CKH\test_file\1.docx")
    print(len(b))
