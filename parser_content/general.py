from tika import parser
from io import BytesIO
from docx import Document
from timeit import default_timer as timer
import re
from parser_file.ragflow_pdf import PlainParser
from nlp import rag_tokenizer, naive_merge, tokenize_table, tokenize_chunks, find_codec, concat_img, \
    naive_merge_docx, tokenize_chunks_docx
from parser_file.ragflow_pdf import RAGFlowPdfParser as  PdfParser
from parser_file.ragflow_docx import RAGFlowDocxParser as DocxParser
from parser_file.ragflow_markdown import RAGFlowMarkdownParser as MarkdownParser
from parser_file.ragflow_txt import RAGFlowTxtParser as TxtParser
from parser_file.ragflow_html import RAGFlowHtmlParser as HtmlParser
from parser_file.ragflow_excel import RAGFlowExcelParser as ExcelParser
from parser_file.ragflow_json import RAGFlowJsonParser as JsonParser
from nlp.token_num import num_tokens_from_string
from PIL import Image
from functools import reduce
from markdown import markdown
from docx.image.exceptions import UnrecognizedImageError, UnexpectedEndOfFileError, InvalidImageStreamError
import os
import settings


class Docx(DocxParser):
    def __init__(self):
        pass

    def get_picture(self, document, paragraph):
        img = paragraph._element.xpath('.//pic:pic')
        if not img:
            return None
        img = img[0]
        embed = img.xpath('.//a:blip/@r:embed')[0]
        related_part = document.part.related_parts[embed]
        try:
            image_blob = related_part.image.blob
        except UnrecognizedImageError:
            print("无法识别图片格式，跳过")
            return None
        except UnexpectedEndOfFileError:
            print("图片EOF错误")
            return None
        except InvalidImageStreamError:
            print("图像输入流错误")
            return None
        try:
            image = Image.open(BytesIO(image_blob)).convert('RGB')
            return image
        except Exception as e:
            return None

    def __clean(self, line):
        line = re.sub(r"\u3000", " ", line).strip()
        return line

    def __call__(self, filename, binary=None, from_page=0, to_page=100000):
        self.doc = Document(
            filename) if not binary else Document(BytesIO(binary))
        pn = 0
        lines = []
        last_image = None
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            if from_page <= pn < to_page:
                if p.text.strip():
                    if p.style and p.style.name == 'Caption':
                        former_image = None
                        if lines and lines[-1][1] and lines[-1][2] != 'Caption':
                            former_image = lines[-1][1].pop()
                        elif last_image:
                            former_image = last_image
                            last_image = None
                        lines.append((self.__clean(p.text), [former_image], p.style.name))
                    else:
                        current_image = self.get_picture(self.doc, p)
                        image_list = [current_image]
                        if last_image:
                            image_list.insert(0, last_image)
                            last_image = None
                        lines.append((self.__clean(p.text), image_list, p.style.name if p.style else ""))
                else:
                    if current_image := self.get_picture(self.doc, p):
                        if lines:
                            lines[-1][1].append(current_image)
                        else:
                            last_image = current_image
            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        new_line = [(line[0], reduce(concat_img, line[1]) if line[1] else None) for line in lines]

        tbls = []
        for tb in self.doc.tables:
            html = "<table>"
            for r in tb.rows:
                html += "<tr>"
                i = 0
                while i < len(r.cells):
                    span = 1
                    c = r.cells[i]
                    for j in range(i + 1, len(r.cells)):
                        if c.text == r.cells[j].text:
                            span += 1
                            i = j
                    i += 1
                    html += f"<td>{c.text}</td>" if span == 1 else f"<td colspan='{span}'>{c.text}</td>"
                html += "</tr>"
            html += "</table>"
            tbls.append(((None, html), ""))
        return new_line, tbls


class Pdf(PdfParser):
    def __call__(self, filename, binary=None, from_page=0,
                 to_page=100000, zoomin=3, callback=None):
        start = timer()
        print("准备启动PDF OCR")
        self.__images__(
            filename if not binary else binary,
            zoomin,
            from_page,
            to_page,
            print
        )
        print("PDF OCR结束")
        print("OCR({}~{}): {}".format(from_page, to_page, timer() - start))

        start = timer()
        self._layouts_rec(zoomin)
        print("布局分析结束")
        self._table_transformer_job(zoomin)
        print("表格分析开始")
        self._text_merge()
        print("文本融合结束")
        tbls = self._extract_table_figure(True, zoomin, True, True)
        # self._naive_vertical_merge()
        self._concat_downward()
        # self._filter_forpages()

        print("布局: {}".format(timer() - start))
        return [(b["text"], self._line_tag(b, zoomin))
                for b in self.boxes], tbls


class Markdown(MarkdownParser):
    def __call__(self, filename, binary=None):
        if binary:
            encoding = find_codec(binary)
            txt = binary.decode(encoding, errors="ignore")
        else:
            with open(filename, "r") as f:
                txt = f.read()
        remainder, tables = self.extract_tables_and_remainder(f'{txt}\n')
        sections = []
        tbls = []
        for sec in remainder.split("\n"):
            if num_tokens_from_string(sec) > 10 * self.chunk_token_num:
                sections.append((sec[:int(len(sec) / 2)], ""))
                sections.append((sec[int(len(sec) / 2):], ""))
            else:
                if sections and sections[-1][0].strip().find("#") == 0:
                    sec_, _ = sections.pop(-1)
                    sections.append((sec_+"\n"+sec, ""))
                else:
                    sections.append((sec, ""))

        for table in tables:
            tbls.append(((None, markdown(table, extensions=['markdown.extensions.tables'])), ""))
        return sections, tbls


def chunk(filename, binary=None, from_page=0, to_page=100000,
          lang="Chinese",chunk_token_num=128, **kwargs):
    """
        Supported file formats are docx, pdf, excel, txt.
        This method apply the naive ways to chunk files.
        Successive text will be sliced into pieces using 'delimiter'.
        Next, these successive pieces are merge into chunks whose token number is no more than 'Max token number'.
    """

    eng = lang.lower() == "english"  # is_english(cks)
    parser_config = kwargs.get(
        "parser_config", {
            "chunk_token_num": chunk_token_num, "delimiter": "\n!?。；！？", "layout_recognize": True})
    doc = {
        "docnm_kwd": os.path.split(filename)[-1],
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", os.path.split(filename)[-1]))
    }
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    res = []
    pdf_parser = None
    if re.search(r"\.docx$", filename, re.IGNORECASE):
        print("开始解析")
        sections, tbls = Docx()(filename, binary)
        res = tokenize_table(tbls, doc, eng)  # just for table

        print("解析结束")
        st = timer()

        chunks, images = naive_merge_docx(
            sections, int(parser_config.get(
                "chunk_token_num", chunk_token_num)), parser_config.get(
                "delimiter", "\n!?。；！？"))

        if kwargs.get("section_only", False):
            return chunks

        res.extend(tokenize_chunks_docx(chunks, doc, eng, images))
        print("naive_merge({}): {}".format(filename, timer() - st))
        return res

    elif re.search(r"\.pdf$", filename, re.IGNORECASE):
        pdf_parser = Pdf(
        ) if parser_config.get("layout_recognize", True) else PlainParser()
        sections, tbls = pdf_parser(filename if not binary else binary,
                                    from_page=from_page, to_page=to_page, print=print)
        res = tokenize_table(tbls, doc, eng)

    elif re.search(r"\.xlsx?$", filename, re.IGNORECASE):
        print("开始解析")
        excel_parser = ExcelParser()
        if parser_config.get("html4excel"):
            sections = [(_, "") for _ in excel_parser.html(binary, 12) if _]
        else:
            sections = [(_, "") for _ in excel_parser(binary) if _]

    elif re.search(r"\.(txt|py|js|java|c|cpp|h|php|go|ts|sh|cs|kt|sql)$", filename, re.IGNORECASE):
        print("开始解析")
        sections = TxtParser()(filename, binary,
                               parser_config.get("chunk_token_num", chunk_token_num),
                               parser_config.get("delimiter", "\n!?;。；！？"))
        print("解析结束")

    elif re.search(r"\.(md|markdown)$", filename, re.IGNORECASE):
        print("开始解析")
        sections, tbls = Markdown(int(parser_config.get("chunk_token_num", chunk_token_num)))(filename, binary)
        res = tokenize_table(tbls, doc, eng)
        print("解析结束")

    elif re.search(r"\.(htm|html)$", filename, re.IGNORECASE):
        print("开始解析")
        sections = HtmlParser()(filename, binary)
        sections = [(_, "") for _ in sections if _]
        print("解析结束")

    elif re.search(r"\.json$", filename, re.IGNORECASE):
        print("开始解析")
        sections = JsonParser(int(parser_config.get("chunk_token_num", chunk_token_num)))(binary)
        sections = [(_, "") for _ in sections if _]
        print("解析结束")

    elif re.search(r"\.doc$", filename, re.IGNORECASE):
        print("正在转换为docx，临时文件在解析后将被自动清理")
        filename=settings.convert_doc_to_docx(filename)
        filename=os.path.join(settings.get_project_base_directory(),filename)
        print("开始解析")
        sections, tbls = Docx()(filename, binary)
        res = tokenize_table(tbls, doc, eng)  # just for table
        print("解析结束")
        st = timer()
        chunks, images = naive_merge_docx(
            sections, int(parser_config.get(
                "chunk_token_num", chunk_token_num)), parser_config.get(
                "delimiter", "\n!?。；！？"))
        if kwargs.get("section_only", False):
            return chunks
        res.extend(tokenize_chunks_docx(chunks, doc, eng, images))
        print("naive_merge({}): {}".format(filename, timer() - st))
        return res

    else:
        raise NotImplementedError(
            "file type not supported yet(pdf, xlsx, doc, docx, txt supported)")

    st = timer()
    chunks = naive_merge(
        sections, int(parser_config.get(
            "chunk_token_num", chunk_token_num)), parser_config.get(
            "delimiter", "\n!?。；！？"))
    if kwargs.get("section_only", False):
        return chunks

    res.extend(tokenize_chunks(chunks, doc, eng, pdf_parser))
    print("naive_merge({}): {}".format(filename, timer() - st))
    return res


if __name__ == "__main__":
    print(chunk(r"E:\Rag-CKH\test_file\2022版煤矿安全规程.docx",chunk_token_num=12))